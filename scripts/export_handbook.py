#!/usr/bin/env python3
"""Build a styled Word handbook from the explicit Zensical navigation."""

from __future__ import annotations

import re
import tomllib
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = ROOT / "deliverables" / "greenhouse-sentinel-handbook.docx"
PRIMARY = RGBColor(7, 94, 84)
ACCENT = RGBColor(180, 83, 9)
TEXT = RGBColor(22, 48, 43)
BODY_FONT = "Carlito"
MONO_FONT = "DejaVu Sans Mono"


def set_font(style, name, size=None):
    """Set a font name consistently for Word and LibreOffice renderers."""
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)


def flatten_nav(items):
    result = []
    for item in items:
        if isinstance(item, str):
            result.append(item)
        else:
            for value in item.values():
                result.extend(flatten_nav(value) if isinstance(value, list) else [value])
    return result


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def repeat_table_header(row):
    """Repeat a table's header row when it continues on another page."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    """Add readable padding around table-cell content."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_column_widths(headers):
    """Use practical widths for the demonstration tables."""
    normalized = tuple(value.lower() for value in headers)
    if normalized == ("feature", "made with", "why it helps"):
        return (1.45, 2.55, 2.5)
    if normalized == ("state", "example condition", "operator expectation"):
        return (1.0, 2.8, 2.7)
    return tuple(6.5 / len(headers) for _ in headers)


def format_callout(paragraph):
    """Give exported admonitions visible separation from surrounding prose."""
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "E8F2F0")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "B45309")
    borders.append(left)
    p_pr.append(borders)


def add_field(paragraph, instruction, fallback="1"):
    """Add a Word field with visible fallback text for headless renderers."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" {instruction} "
    paragraph.add_run()._r.append(code)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    paragraph.add_run(fallback)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def restart_numbering(doc, paragraph):
    """Give a numbered-list block its own sequence beginning at 1."""
    numbering = doc.part.numbering_part.element
    base = next(node for node in numbering.findall(qn("w:num")) if node.get(qn("w:numId")) == "5")
    new_num = deepcopy(base)
    new_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    new_num.set(qn("w:numId"), str(new_id))
    override = OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride"); start.set(qn("w:val"), "1"); override.append(start); new_num.append(override)
    numbering.append(new_num)
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); num_pr.append(ilvl)
    num_id = OxmlElement("w:numId"); num_id.set(qn("w:val"), str(new_id)); num_pr.append(num_id)


def add_inline(paragraph, text):
    text = re.sub(r"\{\s*\.md-button[^}]*\}", "", text).strip()
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`"):
            run = paragraph.add_run(part[1:-1]); run.font.name = MONO_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
        elif part.startswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("["):
            match = re.match(r"\[([^]]+)\]\([^)]+\)", part)
            run = paragraph.add_run(match.group(1) if match else part); run.font.color.rgb = PRIMARY; run.underline = True
        else:
            paragraph.add_run(part)


def add_page(doc, path):
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    table_lines = []
    in_numbered_list = False
    admonition_title = None
    add_space_before_next = False
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(style="Code"); p.add_run(line)
            continue
        if line.startswith("<a id=") or line.startswith("<div") or line.startswith("</div"):
            continue
        if line.startswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in table_lines]
            if len(rows) > 2:
                table = doc.add_table(rows=1, cols=len(rows[0])); table.autofit = False
                widths = table_column_widths(rows[0])
                for i, value in enumerate(rows[0]):
                    table.rows[0].cells[i].text = value; shade(table.rows[0].cells[i], "075E54")
                    for run in table.rows[0].cells[i].paragraphs[0].runs: run.font.color.rgb = RGBColor(255,255,255); run.bold = True
                repeat_table_header(table.rows[0])
                for values in rows[2:]:
                    cells = table.add_row().cells
                    for i, value in enumerate(values):
                        cells[i].text = ""
                        add_inline(cells[i].paragraphs[0], value)
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        cell.width = Inches(widths[i])
                        set_cell_margins(cell)
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.line_spacing = 1.15
                            paragraph.paragraph_format.space_after = Pt(0)
                add_space_before_next = True
            table_lines = []
        if not line:
            in_numbered_list = False
            continue
        if line.startswith("!!!"):
            match = re.match(r'^!!!\s+(\w+)(?:\s+"([^"]+)")?', line)
            kind = match.group(1).title() if match else "Note"
            admonition_title = match.group(2) if match and match.group(2) else kind
            in_numbered_list = False
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            in_numbered_list = False
            level = min(len(heading.group(1)), 3)
            if level == 1 and len(doc.paragraphs) > 4: doc.add_page_break()
            p = doc.add_heading(heading.group(2), level=level)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            if not in_numbered_list: restart_numbering(doc, p)
            in_numbered_list = True
        elif line.startswith("- [ ] "):
            in_numbered_list = False
            p = doc.add_paragraph(style="Checklist")
            marker = p.add_run("[ ] ")
            marker.bold = True
            marker.font.color.rgb = TEXT
            marker.underline = False
            add_inline(p, line[6:])
        elif line.startswith("- "):
            in_numbered_list = False
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, line[2:])
        elif line.startswith("[!") or line.startswith("<"):
            continue
        elif admonition_title is not None:
            in_numbered_list = False
            p = doc.add_paragraph(style="Callout")
            label = p.add_run(f"{admonition_title}. ")
            label.bold = True
            label.font.color.rgb = PRIMARY
            add_inline(p, line)
            format_callout(p)
            admonition_title = None
        else:
            in_numbered_list = False
            p = doc.add_paragraph(); add_inline(p, line)
        if add_space_before_next:
            p.paragraph_format.space_before = Pt(8)
            add_space_before_next = False


def build():
    config = tomllib.loads((ROOT / "zensical.toml").read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.9)
    styles = doc.styles
    normal = styles["Normal"]; set_font(normal, BODY_FONT, 10.5); normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(7); normal.paragraph_format.line_spacing = 1.18
    for name, size in (("Title", 30), ("Heading 1", 21), ("Heading 2", 15), ("Heading 3", 12)):
        style = styles[name]; set_font(style, BODY_FONT, size); style.font.bold = True; style.font.color.rgb = PRIMARY
        style.paragraph_format.line_spacing = 1.05
        style.paragraph_format.space_before = Pt(12); style.paragraph_format.space_after = Pt(6); style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]; set_font(style, BODY_FONT, 10.5); style.font.color.rgb = TEXT
        style.paragraph_format.line_spacing = 1.15; style.paragraph_format.space_after = Pt(2)
    checklist = styles.add_style("Checklist", WD_STYLE_TYPE.PARAGRAPH); set_font(checklist, BODY_FONT, 10.5)
    checklist.font.color.rgb = TEXT; checklist.font.underline = False
    checklist.paragraph_format.left_indent = Inches(0.35); checklist.paragraph_format.first_line_indent = Inches(-0.25)
    checklist.paragraph_format.line_spacing = 1.15; checklist.paragraph_format.space_after = Pt(3)
    callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH); set_font(callout, BODY_FONT, 10)
    callout.font.color.rgb = TEXT; callout.paragraph_format.left_indent = Inches(0.2); callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(8); callout.paragraph_format.space_after = Pt(9); callout.paragraph_format.line_spacing = 1.15
    footer_style = styles["Footer"]; set_font(footer_style, BODY_FONT, 9); footer_style.font.color.rgb = TEXT
    code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH); set_font(code, MONO_FONT, 8.5)
    code.paragraph_format.left_indent = Inches(0.2); code.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph(style="Title"); title.add_run("Greenhouse Sentinel")
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("Product Handbook and Docs-as-Code Portfolio Demonstration"); run.font.size = Pt(15); run.font.color.rgb = ACCENT
    doc.add_paragraph("Fictional product · Real documentation workflow · Katie Kearns · 2026")
    doc.add_paragraph("This public sample contains no customer, proprietary, export-controlled, or classified information.")
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Greenhouse Sentinel · Portfolio demonstration · Katie Kearns · Page ")
    add_field(footer, "PAGE")
    footer.add_run(" of ")
    add_field(footer, "NUMPAGES")
    update_fields = OxmlElement("w:updateFields"); update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)
    for rel in flatten_nav(config["project"]["nav"]): add_page(doc, DOCS / rel)
    OUT.parent.mkdir(exist_ok=True)
    doc.core_properties.title = "Greenhouse Sentinel Handbook"
    doc.core_properties.author = "Katie Kearns"
    doc.core_properties.subject = "Public docs-as-code portfolio demonstration"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__": build()
